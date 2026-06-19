from __future__ import annotations

import os
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "Documentation" / "User Manual"
ASSET_DIR = OUT_DIR / "Assets"
DOCX_PATH = OUT_DIR / "Planetka User Guide.docx"

PAGE_WIDTH_DXA = 9360
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(20, 28, 36)
MUTED = RGBColor(90, 96, 105)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
CALLOUT_FILL = "F4F6F9"
WARNING_FILL = "FFF4CE"
SUCCESS_FILL = "EAF7EF"


def font_path(name: str) -> str | None:
    candidates = [
        f"/System/Library/Fonts/Supplemental/{name}",
        f"/Library/Fonts/{name}",
        f"/System/Library/Fonts/{name}",
    ]
    for candidate in candidates:
        if Path(candidate).exists():
            return candidate
    return None


FONT_REGULAR = font_path("Arial.ttf") or font_path("Arial Unicode.ttf")
FONT_BOLD = font_path("Arial Bold.ttf") or FONT_REGULAR


def pil_font(size: int, bold: bool = False):
    path = FONT_BOLD if bold else FONT_REGULAR
    if path:
        return ImageFont.truetype(path, size)
    return ImageFont.load_default()


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths_dxa):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, width in enumerate(widths_dxa):
            cell = row.cells[idx]
            cell.width = Pt(width / 20)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def set_row_as_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def set_paragraph_border(paragraph, color="2E74B5", size="8", space="6"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = p_bdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        p_bdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color)


def add_run(p, text, bold=False, italic=False, size=None, color=None):
    run = p.add_run(text)
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.bold = bold
    run.italic = italic
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    return run


def add_para(doc, text="", style=None, bold_prefix=None):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.25
    if bold_prefix and text.startswith(bold_prefix):
        add_run(p, bold_prefix, bold=True)
        add_run(p, text[len(bold_prefix):])
    else:
        add_run(p, text)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.25
        if isinstance(item, tuple):
            add_run(p, item[0], bold=True)
            add_run(p, item[1])
        else:
            add_run(p, item)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.25
        add_run(p, item)


def add_callout(doc, title, body, fill=CALLOUT_FILL):
    table = doc.add_table(rows=1, cols=1)
    set_row_as_header(table.rows[0])
    set_table_width(table, [PAGE_WIDTH_DXA - 240])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    add_run(p, title, bold=True, color=DARK_BLUE)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.2
    add_run(p2, body)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_row_as_header(table.rows[0])
    set_table_width(table, widths)
    for idx, header in enumerate(headers):
        cell = table.cell(0, idx)
        set_cell_shading(cell, LIGHT_BLUE)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        add_run(p, header, bold=True, color=DARK_BLUE)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            p = cells[idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.15
            add_run(p, str(value))
    set_table_width(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return table


def rounded_rect(draw, xy, radius, fill, outline=None, width=1):
    draw.rounded_rectangle(xy, radius=radius, fill=fill, outline=outline, width=width)


def draw_label(draw, xy, text, size=18, bold=False, fill=(230, 238, 245)):
    draw.text(xy, text, font=pil_font(size, bold), fill=fill)


def panel_screenshot(name, highlights=(), studio=False, state="before"):
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    path = ASSET_DIR / f"{name}.png"
    w, h = 1100, 1500 if studio else 1260
    img = Image.new("RGB", (w, h), (34, 37, 42))
    draw = ImageDraw.Draw(img)
    # viewport background
    draw.rectangle([0, 0, w, h], fill=(38, 42, 48))
    # sidebar
    panel_x, panel_y, panel_w = 600, 55, 455
    rounded_rect(draw, [panel_x, panel_y, panel_x + panel_w, h - 55], 8, (50, 54, 62), (78, 85, 96), 2)
    draw_label(draw, (panel_x + 22, panel_y + 22), "Planetka by Tomas Griger", 24, True)
    y = panel_y + 76
    boxes = {}

    def box(key, title, lines, height):
        nonlocal y
        x0, x1 = panel_x + 18, panel_x + panel_w - 18
        boxes[key] = (x0, y, x1, y + height)
        rounded_rect(draw, boxes[key], 8, (63, 68, 78), (92, 101, 113), 2)
        draw_label(draw, (x0 + 18, y + 14), title, 20, True)
        line_y = y + 54
        for label, value, tone in lines:
            draw_label(draw, (x0 + 18, line_y), label, 17, False, (194, 202, 212))
            color = (158, 222, 174) if tone == "ok" else (235, 211, 130) if tone == "warn" else (222, 228, 235)
            draw_label(draw, (x0 + 165, line_y), value, 17, True, color)
            line_y += 34
        y += height + 18

    box(
        "data",
        "Planetka Data",
        [("Status", "Connected" if state != "before" else "Ready", "ok" if state != "before" else "neutral"), ("Edition", "Free / Hobby / Pro / Studio", "neutral")],
        145,
    )

    x0, y0, x1, y1 = panel_x + 18, y, panel_x + panel_w - 18, y + 275
    boxes["quality"] = (x0, y0, x1, y1)
    rounded_rect(draw, boxes["quality"], 8, (63, 68, 78), (92, 101, 113), 2)
    draw_label(draw, (x0 + 18, y0 + 14), "Quality Level", 20, True)
    bx_y = y0 + 58
    labels = [("Preview", "quality_preview"), ("Balanced", "quality_balanced"), ("Full", "quality_full")]
    bx_x = x0 + 18
    for label, key in labels:
        bw = 123
        boxes[key] = (bx_x, bx_y, bx_x + bw, bx_y + 42)
        rounded_rect(draw, boxes[key], 6, (78, 84, 96), (115, 125, 141), 1)
        draw_label(draw, (bx_x + 18, bx_y + 11), label, 16, True)
        bx_x += bw + 10
    button_text = "Create New Earth" if state == "before" else "Resolve Planetka"
    boxes["resolve_button"] = (x0 + 18, y0 + 118, x1 - 18, y0 + 168)
    rounded_rect(draw, boxes["resolve_button"], 6, (57, 105, 168), (95, 151, 219), 2)
    draw_label(draw, (x0 + 112, y0 + 132), button_text, 19, True)
    boxes["runtime_status"] = (x0 + 18, y0 + 184, x1 - 18, y0 + 215)
    boxes["download_status"] = (x0 + 18, y0 + 222, x1 - 18, y0 + 253)
    draw_label(draw, (x0 + 18, y0 + 186), "Resolved successfully", 16, False, (220, 228, 235))
    draw_label(draw, (x0 + 18, y0 + 224), "Last download: 42.31 MB, 8 tiles", 16, False, (220, 228, 235))
    y += 293

    x0, y0, x1, y1 = panel_x + 18, y, panel_x + panel_w - 18, y + 150
    boxes["links"] = (x0, y0, x1, y1)
    rounded_rect(draw, boxes["links"], 8, (63, 68, 78), (92, 101, 113), 2)
    draw_label(draw, (x0 + 18, y0 + 14), "Links", 20, True)
    boxes["tutorials"] = (x0 + 18, y0 + 58, x0 + 195, y0 + 104)
    boxes["resources"] = (x0 + 205, y0 + 58, x1 - 18, y0 + 104)
    rounded_rect(draw, boxes["tutorials"], 6, (78, 84, 96), (115, 125, 141), 1)
    rounded_rect(draw, boxes["resources"], 6, (78, 84, 96), (115, 125, 141), 1)
    draw_label(draw, (x0 + 52, y0 + 71), "Tutorials", 17, True)
    draw_label(draw, (x0 + 250, y0 + 71), "Resources", 17, True)
    y += 168

    if studio:
        x0, y0, x1, y1 = panel_x + 18, y, panel_x + panel_w - 18, y + 280
        boxes["animation"] = (x0, y0, x1, y1)
        rounded_rect(draw, boxes["animation"], 8, (63, 68, 78), (92, 101, 113), 2)
        draw_label(draw, (x0 + 18, y0 + 14), "Planetka Animation", 20, True)
        boxes["quick_preview"] = (x0 + 18, y0 + 55, x1 - 18, y0 + 145)
        boxes["final_animation"] = (x0 + 18, y0 + 162, x1 - 18, y0 + 248)
        rounded_rect(draw, boxes["quick_preview"], 6, (72, 78, 90), (115, 125, 141), 1)
        rounded_rect(draw, boxes["final_animation"], 6, (72, 78, 90), (115, 125, 141), 1)
        draw_label(draw, (x0 + 36, y0 + 68), "Quick Preview", 17, True)
        draw_label(draw, (x0 + 36, y0 + 100), "Build Quick Preview     Play / Pause     Clear", 14, False, (220, 228, 235))
        draw_label(draw, (x0 + 36, y0 + 175), "Final Animation Render", 17, True)
        draw_label(draw, (x0 + 36, y0 + 207), "Render Animation", 15, True, (220, 228, 235))

    # mock viewport hint
    draw.ellipse([110, 280, 420, 590], fill=(28, 80, 123), outline=(80, 147, 187), width=3)
    draw.arc([110, 280, 420, 590], 210, 40, fill=(114, 173, 122), width=32)
    draw_label(draw, (118, 620), "Blender Viewport", 24, True, (178, 189, 204))

    for key in highlights:
        if key not in boxes:
            continue
        x0, y0, x1, y1 = boxes[key]
        draw.rounded_rectangle([x0 - 8, y0 - 8, x1 + 8, y1 + 8], radius=10, outline=(255, 210, 68), width=7)

    img.save(path)
    return path


def objects_screenshot():
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    path = ASSET_DIR / "created_objects.png"
    img = Image.new("RGB", (1100, 720), (38, 42, 48))
    draw = ImageDraw.Draw(img)
    rounded_rect(draw, [80, 70, 1020, 650], 10, (50, 54, 62), (82, 90, 103), 2)
    draw_label(draw, (110, 105), "Objects created by Create New Earth", 28, True)
    items = [
        ("Planetka", "Collection"),
        ("Planetka Earth Surface", "Main render surface; transform it like a normal Blender object"),
        ("Planetka Earth Preview", "Low-level preview object used by the Planetka workflow"),
        ("Planetka Sunlight", "Sun light object created at intensity 10, Rotation Y 90 degrees"),
    ]
    y = 180
    for idx, (name, desc) in enumerate(items):
        fill = (63, 68, 78) if idx else (56, 75, 92)
        rounded_rect(draw, [120, y, 980, y + 82], 8, fill, (96, 106, 120), 2)
        draw_label(draw, (150, y + 18), name, 21, True)
        draw_label(draw, (430, y + 20), desc, 17, False, (220, 228, 235))
        y += 105
    draw.rounded_rectangle([115, 270, 985, 570], radius=12, outline=(255, 210, 68), width=6)
    img.save(path)
    return path


def shader_screenshot():
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    path = ASSET_DIR / "shader_controls.png"
    img = Image.new("RGB", (1100, 700), (36, 39, 45))
    draw = ImageDraw.Draw(img)
    draw_label(draw, (70, 55), "Planetka Surface Grading Group", 30, True)
    nodes = [
        (80, 170, 280, 270, "Group Input", "Surface controls"),
        (410, 135, 720, 300, "Surface Grading", "Color, night/day, displacement"),
        (800, 180, 1020, 260, "Material Output", "Surface + Displacement"),
        (80, 410, 395, 520, "Select Sunlight Source Object Here", "Object: Planetka Sunlight"),
    ]
    for x0, y0, x1, y1, title, body in nodes:
        rounded_rect(draw, [x0, y0, x1, y1], 9, (65, 71, 84), (105, 117, 133), 2)
        draw_label(draw, (x0 + 18, y0 + 18), title, 18, True)
        draw_label(draw, (x0 + 18, y0 + 52), body, 15, False, (220, 228, 235))
    draw.line([280, 220, 410, 220], fill=(148, 160, 177), width=5)
    draw.line([720, 220, 800, 220], fill=(148, 160, 177), width=5)
    draw.line([395, 465, 510, 300], fill=(255, 210, 68), width=5)
    draw.rounded_rectangle([70, 400, 405, 530], radius=12, outline=(255, 210, 68), width=6)
    img.save(path)
    return path


def configure_doc(doc: Document):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_run(footer, "Planetka User Guide - 1.0.0", size=9, color=MUTED)


def add_title_page(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(32)
    p.paragraph_format.space_after = Pt(4)
    add_run(p, "Planetka User Guide", size=28, bold=True, color=BLUE)
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(18)
    add_run(p2, "Cinematic Earth Visualisation System for Blender", size=15, color=MUTED)
    p3 = doc.add_paragraph()
    p3.paragraph_format.space_after = Pt(12)
    add_run(p3, "Version 1.0.0 | Generic guide for Free, Hobby, Pro, and Studio editions", size=11, bold=True)
    rule = doc.add_paragraph()
    set_paragraph_border(rule, color="2E74B5", size="12", space="8")
    add_callout(
        doc,
        "Purpose",
        "This guide explains the Planetka side panel, the Create New Earth and Resolve Planetka workflows, quality choices, object handling, supported cameras, licensing differences, and basic troubleshooting.",
        fill=SUCCESS_FILL,
    )
    add_para(
        doc,
        "Planetka is intentionally focused: it creates a Planetka Earth setup and streams the texture data needed for the current camera or active viewport. Scene composition, camera animation, additional lights, atmosphere, clouds, and other creative elements remain under normal Blender control.",
    )


def add_image(doc, path, caption, width=5.9):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    inline_shape = run.add_picture(str(path), width=Inches(width))
    doc_pr = inline_shape._inline.docPr
    doc_pr.set("descr", caption)
    doc_pr.set("title", Path(path).stem.replace("_", " "))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(10)
    add_run(cap, caption, size=9, italic=True, color=MUTED)


def build_manual():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    shots = {
        "data": panel_screenshot("side_panel_data", ["data"], state="after"),
        "create": panel_screenshot("side_panel_create", ["resolve_button"], state="before"),
        "quality": panel_screenshot("side_panel_quality", ["quality", "quality_preview", "quality_balanced", "quality_full"], state="after"),
        "progress": panel_screenshot("side_panel_progress", ["runtime_status", "download_status"], state="after"),
        "links": panel_screenshot("side_panel_links", ["links", "tutorials", "resources"], state="after"),
        "studio": panel_screenshot("side_panel_studio_animation", ["animation", "quick_preview", "final_animation"], studio=True, state="after"),
        "objects": objects_screenshot(),
        "shader": shader_screenshot(),
    }

    doc = Document()
    configure_doc(doc)
    add_title_page(doc)

    doc.add_heading("1. What Planetka Does", level=1)
    add_para(doc, "Planetka creates and resolves a cinematic Earth surface inside Blender. It supplies the Earth mesh, preview object, sunlight object, locked Planetka materials, and a manual Resolve workflow that streams visible texture data from Planetka Cloud.")
    add_bullets(doc, [
        ("Create New Earth: ", "creates the Planetka collection, Earth surface, preview object, sunlight object, and material setup."),
        ("Resolve Planetka: ", "streams and applies texture data for the current camera view or active viewport."),
        ("Quality Level: ", "controls the downloaded texture size used for Resolve."),
        ("Studio animation tools: ", "Studio-only tools for Quick Preview and Final Animation Render."),
    ])
    add_callout(doc, "Important", "Planetka is not a full scene builder. You use normal Blender tools for your camera, render settings, object transform, scene lighting, atmosphere, clouds, compositing, and animation paths.")

    doc.add_heading("2. Editions and Licensing", level=1)
    add_table(
        doc,
        ["Edition", "Package", "Texture access", "Animation panel", "Licence / support"],
        [
            ["Free", "Unsigned ZIP", "1000K texture access, capped at d004 or coarser", "No", "Personal use only; community support; best-effort data access"],
            ["Hobby", "Signed ZIP", "Full Planetka texture range", "No", "Personal use only; community support; best-effort data access"],
            ["Pro", "Signed ZIP", "Full Planetka texture range", "No", "Commercial use; direct support; guaranteed data access"],
            ["Studio", "Signed ZIP", "Full Planetka texture range", "Yes", "Commercial use; multiple seats; direct support; guaranteed data access"],
        ],
        [1200, 1450, 2500, 1300, 2910],
    )
    add_para(doc, "The editions are separate ZIP files, but they are editions of the same Blender add-on. Install only one Planetka edition at a time. Installing a different edition replaces the previous one.")
    add_para(doc, "Free is limited by the backend and by the add-on to d004 or coarser texture data. If a Free user moves very close to Earth, Planetka requests the closest allowed d004-level data instead of showing a d001/d002 access error.")
    add_para(doc, "Hobby and Pro have identical add-on functionality. They differ by licence, support level, and data-access guarantee. Studio adds the animation panel.")
    add_callout(doc, "Hosted data terms", "The Blender add-on source code is GPLv3. Planetka Cloud, hosted texture data, cached texture data, edition signatures, account access, API access, service availability, and the processed Planetka texture database are governed by Planetka Terms, Data Licence, and Fair Usage Policy.")

    doc.add_heading("3. Compatibility", level=1)
    add_para(doc, "Planetka is developed primarily for Blender 5.0 and targets Blender 4.5 LTS and newer supported Blender versions.")
    add_table(
        doc,
        ["Area", "Supported", "Notes"],
        [
            ["Blender", "Blender 4.5 LTS and newer supported versions", "Use Blender 5.0 or newer where possible."],
            ["Internet", "Required", "Create Earth and Resolve require Planetka Cloud availability."],
            ["Cameras", "Perspective and Orthographic cameras", "Resolve uses the active scene camera when you are in camera view."],
            ["Active View", "Perspective and Orthographic viewport views", "Resolve can use the current 3D View instead of the camera."],
            ["Panorama", "Equirectangular panorama only, with limits", "Other panoramic and fisheye projections are not currently supported."],
            ["GPU / memory", "High-resolution texture capable GPU recommended", "Full quality can use more memory and texture slots."],
        ],
        [1500, 2700, 5160],
    )

    doc.add_heading("4. The Planetka Side Panel", level=1)
    add_image(doc, shots["data"], "Planetka Data status and edition fields.")
    add_para(doc, "Open the Planetka side panel in the 3D View sidebar. The main panel is named Planetka by Tomas Griger.")
    add_bullets(doc, [
        ("Status: ", "shows Planetka Data state. Typical values are Ready before connection, Connected after a session is established, or a backend message such as maintenance or temporary unavailability."),
        ("Edition: ", "shows the installed edition detected from the local Planetka package."),
        ("Quality Level: ", "selects Preview, Balanced, or Full for the next Resolve."),
        ("Create New Earth / Resolve Planetka: ", "the same workflow button changes label after Earth is created."),
        ("Status and progress lines: ", "show current operation and download amount without shifting the panel layout."),
        ("Links: ", "open tutorials and Planetka Resources in your browser."),
    ])

    doc.add_heading("5. First Start Guide", level=1)
    add_numbered(doc, [
        "Install the Planetka ZIP in Blender through Preferences > Add-ons or Extensions, depending on your Blender version.",
        "Open a 3D View and show the sidebar with N.",
        "Select the Planetka tab.",
        "Set your scene camera or prepare the 3D View you want to resolve.",
        "Press Create New Earth.",
        "Wait until the operation finishes and the status line reports that Planetka is resolved or ready.",
        "Choose a quality level and press Resolve Planetka whenever you want to update the visible texture data.",
    ])
    add_callout(doc, "Connection timing", "Planetka connects to Planetka Cloud when you press Create New Earth or Resolve Planetka. Opening the side panel alone should not start a cloud session.")

    doc.add_heading("6. Create New Earth", level=1)
    add_image(doc, shots["create"], "Before Earth exists, the main action button is Create New Earth.")
    add_para(doc, "Create New Earth prepares the Planetka scene objects and loads initial preview data for the current camera or active viewport.")
    add_table(
        doc,
        ["Created item", "Purpose"],
        [
            ["Planetka collection", "Holds the Planetka-created objects."],
            ["Planetka Earth Surface", "Main Earth object used for rendering and resolving."],
            ["Planetka Earth Preview", "Preview object used by Planetka's workflow."],
            ["Planetka Sunlight", "Sun light object with intensity 10 and Rotation Y set to 90 degrees."],
        ],
        [2300, 7060],
    )
    add_image(doc, shots["objects"], "Create New Earth creates only the Planetka Earth objects and Planetka Sunlight.")
    add_para(doc, "Use your own Blender camera, camera animation, constraints, and scene composition.")
    add_para(doc, "Add atmosphere, cloud, or other scene elements separately as normal Blender scene elements.")

    doc.add_heading("7. Moving and Scaling Earth", level=1)
    add_para(doc, "Treat Planetka Earth Surface as a normal Blender object. You can move, rotate, scale, or apply transforms to place Earth in your scene.")
    add_bullets(doc, [
        "Resolve Planetka should preserve the Earth object's transform, dimensions, scale, material displacement setting, and user-facing material settings.",
        "If you change Dimensions, Blender may change Scale automatically. That is normal Blender behavior.",
        "If you apply scale, Planetka updates its internal displacement scaling so the apparent displacement remains consistent.",
        "If a Resolve appears to do nothing after major scaling, check that your camera or active view is outside the Earth and looking at the surface.",
    ])
    add_callout(doc, "Recommended workflow", "Set the Earth size and position first, then place the camera. After large scale changes, press Resolve Planetka again so the texture selection matches the new camera distance.")

    doc.add_heading("8. Planetka Sunlight and Shader Controls", level=1)
    add_image(doc, shots["shader"], "The sunlight source is selected inside the material shader node setup.")
    add_para(doc, "Create New Earth creates Planetka Sunlight and assigns it to the shader control node named Select Sunlight Source Object Here. This object is used by the Planetka surface shader to drive the day/night transition.")
    add_bullets(doc, [
        "You can rotate or replace the sunlight object if your scene requires a different light direction.",
        "If the day/night transition appears wrong, check that the shader node still points to the intended sunlight object.",
        "Most Earth visual controls are available inside the Planetka Surface Grading Group, not as separate side-panel settings.",
        "The Planetka Earth Material and Planetka Preview Material use the same default values for their shared surface grading group.",
    ])

    doc.add_heading("9. Resolve Planetka", level=1)
    add_para(doc, "Resolve Planetka reads the active camera or active viewport, determines which Earth tiles are visible, streams missing data from Planetka Cloud or local cache, and rebuilds the Planetka Earth Surface mesh/material assignment.")
    add_image(doc, shots["quality"], "Quality Level controls the texture size used by the next Resolve.")
    add_table(
        doc,
        ["Quality", "Meaning", "Use when"],
        [
            ["Preview", "Downloaded textures are 1/4 of the edge size of Full resolution textures, making them 1/16 of the pixel size.", "Fast look development, blocking, camera setup, early tests."],
            ["Balanced", "Downloaded textures are 1/2 of the edge size of Full resolution textures, making them 1/4 of the pixel size.", "General iteration when Preview is too soft but Full is unnecessary."],
            ["Full", "Makes sure at least one pixel from the source texture is used for every pixel in the final render if proximity to Earth allows.", "Final stills or high-quality checks."],
        ],
        [1300, 5000, 3060],
    )
    add_para(doc, "Changing Quality Level does not change the Earth immediately. Press Resolve Planetka after changing quality.")
    add_image(doc, shots["progress"], "The upper line shows operation status; the lower line shows download amount or the last completed download.")

    doc.add_heading("10. Camera View and Active View Resolving", level=1)
    add_para(doc, "Planetka can resolve for the scene camera or for the current 3D View.")
    add_bullets(doc, [
        ("Camera view: ", "when you resolve from camera view, Planetka uses the active scene camera."),
        ("Active View: ", "when you resolve from a normal 3D viewport, Planetka uses the current viewport direction, zoom, and projection."),
        ("Perspective cameras: ", "supported."),
        ("Orthographic cameras: ", "supported."),
        ("Equirectangular panorama cameras: ", "supported with tile-budget limits."),
        ("Other panoramic and fisheye projections: ", "not currently supported."),
    ])
    add_callout(doc, "Practical note", "For final work, resolve from the view you will actually render. If you change camera position, focal length, resolution, or Earth size, resolve again.")

    doc.add_heading("11. Links", level=1)
    add_image(doc, shots["links"], "Tutorials and Resources open in your browser.")
    add_bullets(doc, [
        ("Tutorials: ", "opens Planetka tutorial videos."),
        ("Resources: ", "opens Planetka resources, where related scene elements can be provided."),
    ])

    doc.add_heading("12. Studio Animation Panel", level=1)
    add_para(doc, "The animation panel appears only in the Studio edition. Free, Hobby, and Pro do not include the animation panel.")
    add_image(doc, shots["studio"], "Studio-only animation panel with Quick Preview and Final Animation Render.", width=4.6)
    add_bullets(doc, [
        ("Build Quick Preview: ", "downloads Preview-quality data for animation segments and builds preview segment meshes/materials."),
        ("Play / Pause: ", "plays or pauses the prepared preview animation."),
        ("Clear: ", "removes prepared animation preview assets."),
        ("Render Animation: ", "runs Final Animation Render for the active timeline setup."),
        ("Stop: ", "appears while Final Animation Render is active and stops the active render workflow."),
    ])
    add_callout(doc, "Studio-only backend access", "Final Animation Render tile sessions are allowed only for Studio. Free, Hobby, and Pro cannot start Final Animation Render sessions from the backend.")

    doc.add_heading("13. Cache, Cloud Data, and Fair Usage", level=1)
    add_para(doc, "Planetka streams required data from Planetka Cloud and keeps a working cache for recent Resolves. It is not a raw dataset download product.")
    add_bullets(doc, [
        "If a needed file does not exist in the cloud database, Planetka may use bundled fallback images for cases such as ocean or polar fallback data.",
        "Normal Earth texture tiles are pulled from Cloudflare unless already available in cache.",
        "Data usage per installation is reviewed regularly. Excessive or suspicious use may be throttled or blocked.",
        "For unrestricted access, API access, or custom data arrangements, contact info@planetka.io.",
    ])

    doc.add_heading("14. Troubleshooting", level=1)
    add_table(
        doc,
        ["Issue", "Likely cause", "What to do"],
        [
            ["Status stays Ready before first use", "Planetka has not connected yet.", "Press Create New Earth or Resolve Planetka. The side panel alone does not connect."],
            ["No internet at startup", "Cloud session cannot be created.", "Reconnect internet and press Create New Earth or Resolve again."],
            ["Internet drops during Resolve", "Download was interrupted.", "Reconnect internet and retry Resolve. Cached files may reduce the next download."],
            ["Cloud busy / maintenance status", "Backend service status message.", "Wait and retry later. If a link is shown, open it for details."],
            ["Pink material", "Missing image or failed texture assignment.", "Retry Resolve. If it persists, check internet and report the file/camera state."],
            ["Black render", "Camera on night side, lighting, exposure, or missing surface visibility.", "Check Planetka Sunlight direction, camera position, render exposure, and material assignment."],
            ["Earth changes size after edits", "Object transforms or Apply Scale changed local mesh/scale relationship.", "Planetka should preserve transforms on Resolve. If needed, set dimensions again and Resolve."],
            ["Resolve seems to do nothing", "Camera/Active View may not see the Earth or may be inside the surface.", "Move camera outside Earth, aim at the surface, then Resolve."],
            ["Missing detail close to Earth in Free", "Free is capped at d004 or coarser.", "This is expected. Hobby, Pro, and Studio have full texture range."],
            ["Animation panel missing", "Installed edition is Free, Hobby, or Pro.", "Only Studio includes the animation panel."],
            ["Fisheye/panoramic view looks wrong", "Projection is not supported.", "Use Perspective, Orthographic, or Equirectangular panorama with limits."],
        ],
        [2300, 3100, 3960],
    )

    doc.add_heading("15. Support Checklist", level=1)
    add_para(doc, "When reporting an issue, include enough scene context to reproduce it.")
    add_bullets(doc, [
        "Blender version and operating system.",
        "Planetka edition and version.",
        "Whether you resolved from Camera View or Active View.",
        "Camera type and render resolution.",
        "Quality level selected.",
        "Whether the Earth object was moved, scaled, or had transforms applied.",
        "Exact status/error text from the Planetka panel.",
        "A small test .blend file if the issue is visual or camera-dependent.",
    ])

    doc.add_heading("16. Quick Reference", level=1)
    add_table(
        doc,
        ["Task", "Action"],
        [
            ["Create Planetka Earth", "Press Create New Earth."],
            ["Update visible texture data", "Choose quality, then press Resolve Planetka."],
            ["Resolve current camera", "Use camera view or set active scene camera, then Resolve."],
            ["Resolve current viewport", "Use Active View in the 3D View, then Resolve."],
            ["Move or resize Earth", "Transform Planetka Earth Surface normally in Blender."],
            ["Change day/night direction", "Adjust Planetka Sunlight or the shader's sunlight object reference."],
            ["Open learning resources", "Use Tutorials or Resources under Links."],
            ["Prepare animation tools", "Use Studio edition only."],
        ],
        [3000, 6360],
    )

    doc.save(DOCX_PATH)
    return DOCX_PATH


if __name__ == "__main__":
    print(build_manual())
